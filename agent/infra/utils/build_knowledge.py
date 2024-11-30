import re

import pymupdf4llm
from docx import Document
from llama_ocr import ocr
import pytesseract
from PIL import Image
from pytesseract import Output
import cv2
import numpy as np

def preprocess_image(image_path):
    img = cv2.imread(image_path)
    # 转灰度图
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # 二值化处理

    _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # 降噪

    denoised = cv2.fastNlMeansDenoising(binary)

    return denoised
#doc = Document("docs/冰雪极地_14306704.docx")

#pdf = pymupdf4llm.to_markdown("docs/冰雪极地_14306704.pdf", write_images=True, image_path="./docs", image_format="png")

#with open("test.md", "w") as f:
#    f.write(pdf)

#rmd = ocr("docs/冰雪极地_14306704.pdf-58-0.png",)
#img = Image.open('docs/冰雪极地_14306704.pdf-10-0.png')
img = preprocess_image('docs/冰雪极地_14306704.pdf-10-0.png')
text = pytesseract.image_to_string(img, lang="chi_sim", output_type=Output.DICT)

print(text)
'''
for table in doc.tables:
    print(table)

for para in doc.paragraphs:
    para.
    if para is None or len(para.text) == 0:
        continue
    print(para.text)
'''